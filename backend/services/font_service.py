import os
import subprocess
import requests
from typing import Set, List
from docx import Document


class FontService:
    """Service for managing fonts needed for PDF conversion."""
    
    def __init__(self):
        self.fonts_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'fonts')
        os.makedirs(self.fonts_dir, exist_ok=True)
    
    def get_fonts_from_docx(self, docx_path: str) -> Set[str]:
        """Extract all font names used in a DOCX file."""
        try:
            doc = Document(docx_path)
            fonts = set()
            
            # Get fonts from paragraphs
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font.name:
                        fonts.add(run.font.name)
            
            # Get fonts from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.font.name:
                                    fonts.add(run.font.name)
            
            return fonts
        except Exception as e:
            print(f"   ⚠️  Error extracting fonts: {e}")
            return set()
    
    def is_font_available(self, font_name: str) -> bool:
        """Check if a font is available on the system."""
        try:
            result = subprocess.run(
                ['fc-list', ':', 'family'],
                capture_output=True,
                timeout=5,
                text=True
            )
            
            if result.returncode == 0:
                available_fonts = result.stdout.lower()
                return font_name.lower() in available_fonts
            
            return False
        except Exception:
            return False
    
    def download_google_font(self, font_name: str) -> bool:
        """
        Download a font from Google Fonts to local fonts directory.
        Returns True if successful.
        """
        try:
            # Google Fonts API to get font files
            api_url = f"https://fonts.google.com/download?family={font_name.replace(' ', '+')}"
            
            # Common Google Fonts direct download (using webfonts helper)
            # This is a more reliable method
            font_name_normalized = font_name.lower().replace(' ', '-')
            
            # Try to download from a CDN or direct source
            # For now, we'll use a simpler approach - check if it's a common Google Font
            common_google_fonts = {
                'montserrat': 'https://github.com/JulietaUla/Montserrat/archive/refs/heads/master.zip',
                'roboto': 'https://github.com/google/roboto/archive/refs/heads/main.zip',
                'open sans': 'https://github.com/googlefonts/opensans/archive/refs/heads/main.zip',
                'lato': 'https://github.com/latofonts/lato-source/archive/refs/heads/master.zip',
            }
            
            font_key = font_name.lower()
            if font_key in common_google_fonts:
                print(f"   📥 Font '{font_name}' can be downloaded from Google Fonts")
                print(f"   💡 To install: Download from {common_google_fonts[font_key]}")
                # For now, we'll note it but not auto-download (would need unzipping, etc.)
                return False
            
            return False
            
        except Exception as e:
            print(f"   ⚠️  Error downloading font '{font_name}': {e}")
            return False
    
    def get_font_substitution(self, font_name: str) -> str:
        """
        Determine what font LibreOffice will substitute for a missing font.
        Returns the substitute font name.
        """
        try:
            # Check LibreOffice's font substitution table
            # Common substitutions LibreOffice uses:
            serif_fonts = ['times new roman', 'liberation serif', 'georgia']
            sans_serif_fonts = ['arial', 'liberation sans', 'helvetica', 'dejavu sans']
            monospace_fonts = ['courier new', 'liberation mono', 'dejavu sans mono']
            
            font_lower = font_name.lower()
            
            # Check what's actually available as fallback
            result = subprocess.run(
                ['fc-match', font_name],
                capture_output=True,
                timeout=5,
                text=True
            )
            
            if result.returncode == 0:
                # Parse output like: "DejaVuSans.ttf: "DejaVu Sans" "Book""
                match_line = result.stdout.strip()
                if ':' in match_line:
                    substitute = match_line.split(':')[1].strip().split('"')[1]
                    return substitute
            
            # Fallback: guess based on font characteristics
            if any(keyword in font_lower for keyword in ['serif', 'times', 'georgia']):
                return 'Liberation Serif'
            elif any(keyword in font_lower for keyword in ['mono', 'courier', 'console']):
                return 'Liberation Mono'
            else:
                return 'Liberation Sans'  # Default sans-serif
                
        except Exception:
            return 'Unknown (likely Liberation Sans)'
    
    def ensure_fonts_available(self, docx_path: str) -> dict:
        """
        Check which fonts are needed but not available.
        Returns dict with font status and substitutions.
        
        Returns:
        {
            'fonts_detected': ['Montserrat', 'Arial'],
            'missing_fonts': [
                {
                    'font': 'Montserrat',
                    'substitute': 'Liberation Sans',
                    'install_command': 'sudo pacman -S ttf-montserrat'
                }
            ],
            'all_available': False
        }
        """
        fonts_needed = self.get_fonts_from_docx(docx_path)
        result = {
            'fonts_detected': sorted(list(fonts_needed)),
            'missing_fonts': [],
            'all_available': True
        }
        
        if not fonts_needed:
            return result
        
        print(f"   🔤 Fonts detected: {', '.join(sorted(fonts_needed))}")
        
        for font in fonts_needed:
            if not self.is_font_available(font):
                result['all_available'] = False
                substitute = self.get_font_substitution(font)
                
                missing_info = {
                    'font': font,
                    'substitute': substitute,
                    'install_command': f"sudo pacman -S ttf-{font.lower().replace(' ', '-')}"
                }
                result['missing_fonts'].append(missing_info)
                
                print(f"   ⚠️  Font '{font}' not installed")
                print(f"       → Will be substituted with: {substitute}")
                print(f"       → Install with: {missing_info['install_command']}")
        
        if result['all_available']:
            print(f"   ✅ All fonts available")
        
        return result
