from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from typing import Dict, Any, List
import json
import re
import zipfile
import shutil
import os
from lxml import etree
import tempfile
import math
import subprocess


class DocumentService:
    """Service for document parsing and generation."""
    
    @staticmethod
    def get_page_count(docx_path: str) -> int:
        """
        Get the page count of a DOCX file using python-docx.
        This is an approximation based on the document structure.
        """
        try:
            doc = Document(docx_path)
            # Check if there are explicit page breaks or section properties that indicate page count
            # For more accurate counting, we'd use Word's API or convert to PDF
            # For now, we'll use a simple heuristic or return 0 to trigger calculation
            
            # Try to extract page count from document properties
            # DOCX doesn't store page count reliably, so this is approximate
            # A better approach would be to convert to PDF and count pages
            # For now, return 0 to indicate we need to measure differently
            return 0
        except Exception as e:
            print(f"Error getting page count: {e}")
            return 0
    
    @staticmethod
    def estimate_page_count_from_content(docx_path: str) -> int:
        """
        Estimate page count based on content length and formatting.
        This is a rough approximation.
        """
        try:
            doc = Document(docx_path)
            
            # Count total characters and lines
            total_chars = 0
            total_paragraphs = 0
            
            for para in doc.paragraphs:
                if para.text.strip():
                    total_chars += len(para.text)
                    total_paragraphs += 1
            
            # Rough estimate: ~3000 characters per page with standard formatting
            # Adjust based on paragraph count (more paragraphs = more spacing)
            estimated_pages = max(1, int((total_chars / 3000) + (total_paragraphs / 40)))
            
            return estimated_pages
        except Exception as e:
            print(f"Error estimating page count: {e}")
            return 1
    
    @staticmethod
    def sanitize_text(text: str, strip_markdown: bool = True) -> str:
        """
        Sanitize text to be XML/DOCX compatible by removing control characters.
        Keeps only valid Unicode characters and common whitespace.
        """
        if not isinstance(text, str):
            return str(text)
        
        # Remove control characters except tab, newline, and carriage return
        # XML 1.0 valid characters: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD]
        cleaned = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
        
        # Also remove any null bytes that might have slipped through
        cleaned = cleaned.replace('\x00', '')
        
        if strip_markdown:
            # Remove literal markdown bold/italic asterisks
            cleaned = cleaned.replace('**', '').replace('*', '')
        
        return cleaned
    
    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a DOCX file and extract structured content.
        """
        doc = Document(file_path)
        
        resume_data = {
            "sections": [],
            "full_text": [],
            "formatting": {
                "font": None,
                "has_tables": False
            }
        }
        
        current_section = None
        
        # Only process main document paragraphs (headers/footers are preserved automatically)
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            resume_data["full_text"].append(text)
            
            # Detect section headers (usually bold or larger font)
            is_header = False
            if para.runs:
                first_run = para.runs[0]
                if first_run.bold or (first_run.font.size and first_run.font.size >= Pt(14)):
                    is_header = True
            
            if is_header:
                # Start new section
                if current_section:
                    resume_data["sections"].append(current_section)
                current_section = {
                    "title": text,
                    "content": []
                }
            else:
                # Add to current section
                if current_section:
                    current_section["content"].append(text)
                else:
                    # Create a default section if none exists
                    current_section = {
                        "title": "Header",
                        "content": [text]
                    }
        
        # Add last section
        if current_section:
            resume_data["sections"].append(current_section)
        
        # Check for tables
        if doc.tables:
            resume_data["formatting"]["has_tables"] = True
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                resume_data["sections"].append({
                    "title": "Table",
                    "type": "table",
                    "content": table_data
                })
        
        return resume_data
    
    def create_docx_with_xml_preservation(self, original_file_path: str, resume_data: Dict[str, Any], output_path: str):
        """
        Create a DOCX file using XML-level manipulation to preserve ALL elements including shapes and backgrounds.
        This works by manipulating the document.xml directly without using python-docx's save method.
        """
        # DOCX files are ZIP archives containing XML files
        # Extract, modify XML, and repackage
        
        # Create a temporary directory for extraction
        with tempfile.TemporaryDirectory() as temp_dir:
            # Extract the original DOCX
            with zipfile.ZipFile(original_file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # Path to the main document XML
            document_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
            
            # Parse the XML
            parser = etree.XMLParser(remove_blank_text=False)
            tree = etree.parse(document_xml_path, parser)
            root = tree.getroot()
            
            # Define namespaces
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
                'wpc': 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas',
                'cx': 'http://schemas.microsoft.com/office/drawing/2014/chartex',
                'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
                'o': 'urn:schemas-microsoft-com:office:office',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                'v': 'urn:schemas-microsoft-com:vml',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
                'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006'
            }
            
            # Get full text from resume data
            full_text_items = resume_data.get("full_text", [])
            
            if not full_text_items:
                # Fallback
                full_text_items = []
                for section in resume_data.get("sections", []):
                    if section.get("type") == "table":
                        continue
                    if section.get("title"):
                        full_text_items.append(section["title"])
                    for item in section.get("content", []):
                        if item:
                            full_text_items.append(str(item))
            
            # Collect all paragraph elements from main document ONLY
            # We explicitly SKIP headers/footers to preserve contact info/name
            all_paragraphs = []
            
            # Process main document paragraphs
            # STRICTLY match python-docx behavior: Only direct children of w:body
            body = root.find('w:body', namespaces)
            if body is not None:
                paragraphs = body.findall('w:p', namespaces)
            else:
                paragraphs = []

            for para in paragraphs:
                text_content = ''
                for t_elem in para.findall('.//w:t', namespaces):
                    if t_elem.text:
                        text_content += t_elem.text
                
                if text_content.strip():
                    all_paragraphs.append((para, document_xml_path, text_content))
            
            # Check if counts match
            if len(all_paragraphs) != len(full_text_items):
                print(f"⚠️  Warning: XML paragraph count mismatch!")
                print(f"   Original: {len(all_paragraphs)} paragraphs (body only)")
                print(f"   AI Generated: {len(full_text_items)} text items")
                print(f"   Proceeding with XML preservation anyway (safest option for shapes).")
                # DO NOT FALL BACK - The fallback method destroys shapes
                # self.create_docx_preserve_formatting(original_file_path, resume_data, output_path)
                      # Track which files have been modified
            modified_files = set()
            
            # Update each paragraph's text with preservation
            from collections import deque
            paragraphs_deque = deque(all_paragraphs)
            text_items_deque = deque(full_text_items)
            
            while paragraphs_deque and text_items_deque:
                (para, para_file, old_text) = paragraphs_deque.popleft()
                new_text = text_items_deque.popleft()
                
                # Use raw text (including **) for splitting markers
                raw_new_text = str(new_text)
                
                # Find all text runs in this paragraph
                text_runs = para.findall('.//w:r', namespaces)
                if not text_runs: continue

                # Identify text-bearing runs, skipping those with drawings/picts/etc.
                text_bearing_runs = []
                for run in text_runs:
                    # Comprehensive check for non-text artifacts to preserve
                    has_artifact = (run.find('.//w:drawing', namespaces) is not None or 
                                  run.find('.//w:pict', namespaces) is not None or 
                                  run.find('.//w:object', namespaces) is not None or
                                  run.find('.//v:rect', namespaces) is not None or
                                  run.find('.//v:shape', namespaces) is not None or
                                  run.find('.//wps:wsp', namespaces) is not None or
                                  run.find('.//mc:AlternateContent', namespaces) is not None)
                                  
                    t_elems = run.findall('.//w:t', namespaces)
                    if not has_artifact and t_elems:
                        text_bearing_runs.append((run, t_elems))

                if not text_bearing_runs: continue
                
                # BOLDING LOGIC: Handle **bold** markers from AI output
                # Check for Markdown Bold markers
                md_bold_matches = list(re.finditer(r'\*\*(.*?)\*\*', raw_new_text))
                
                if md_bold_matches:
                    # Use markdown segments to update runs
                    segments = []
                    last_idx = 0
                    for match in md_bold_matches:
                        # Add text before the bold part
                        if match.start() > last_idx:
                            segments.append((raw_new_text[last_idx:match.start()], False))
                        # Add the bold part (content between **)
                        segments.append((match.group(1), True))
                        last_idx = match.end()
                    # Add remaining text
                    if last_idx < len(raw_new_text):
                        segments.append((raw_new_text[last_idx:], False))
                    
                    # Map segments onto available runs
                    for seg_idx, (seg_text, is_bold) in enumerate(segments):
                        if seg_idx < len(text_bearing_runs):
                            run, t_elems = text_bearing_runs[seg_idx]
                            # Update text
                            t_elems[0].text = self.sanitize_text(seg_text, strip_markdown=True)
                            for t in t_elems[1:]: t.getparent().remove(t)
                            
                            # Update bold property strictly on this run
                            rPr = run.find('w:rPr', namespaces)
                            if rPr is None:
                                rPr = etree.SubElement(run, f"{{{namespaces['w']}}}rPr")
                            
                            b_elem = rPr.find('w:b', namespaces)
                            if is_bold:
                                if b_elem is None:
                                    # Add if missing
                                    etree.SubElement(rPr, f"{{{namespaces['w']}}}b")
                                else:
                                    # Ensure it's not disabled
                                    b_elem.attrib.clear() # Clear val="0" if it exists
                            else:
                                if b_elem is not None:
                                    # Explicitly disable bold
                                    b_elem.set(f"{{{namespaces['w']}}}val", "0")
                        else:
                            # If we have more segments than runs, append to the last run
                            # (A more advanced implementation would create new runs)
                            last_run, last_ts = text_bearing_runs[-1]
                            last_ts[0].text += self.sanitize_text(seg_text, strip_markdown=True)
                            
                    # Clear any remaining runs that weren't used
                    for i in range(len(segments), len(text_bearing_runs)):
                        r, ts = text_bearing_runs[i]
                        for t in ts: 
                            if t.getparent() == r: r.remove(t)
                        # Remove the run if it's basically empty
                        if not any(child.tag in [f"{{{namespaces['w']}}}drawing", f"{{{namespaces['w']}}}pict", f"{{{namespaces['mc']}}}AlternateContent", f"{{{namespaces['w']}}}t"] for child in r):
                            if r.getparent() is not None:
                                r.getparent().remove(r)
                else:
                    # Fallback to colon pattern if no markdown markers but colon exists
                    has_colon = ':' in raw_new_text and len(text_bearing_runs) >= 2
                    sanitized_full = self.sanitize_text(raw_new_text, strip_markdown=True)
                    
                    if has_colon:
                        label, content = sanitized_full.split(':', 1)
                        # Label run
                        r1, t1s = text_bearing_runs[0]
                        t1s[0].text = label + ":"
                        for t in t1s[1:]: t.getparent().remove(t)
                        # Value run
                        r2, t2s = text_bearing_runs[1]
                        t2s[0].text = content
                        for t in t2s[1:]: t.getparent().remove(t)
                        
                        # Set bold for label, normal for value
                        for i, (r, ts) in enumerate(text_bearing_runs):
                            rPr = r.find('w:rPr', namespaces)
                            if rPr is None: rPr = etree.SubElement(r, f"{{{namespaces['w']}}}rPr")
                            b_elem = rPr.find('w:b', namespaces)
                            if i == 0: # label
                                if b_elem is None: etree.SubElement(rPr, f"{{{namespaces['w']}}}b")
                                else: b_elem.attrib.clear()
                            elif i == 1: # content
                                if b_elem is not None: b_elem.set(f"{{{namespaces['w']}}}val", "0")
                            else: # clear others
                                for t in ts: 
                                    if t.getparent() == r: r.remove(t)
                                if not any(child.tag in [f"{{{namespaces['w']}}}drawing", f"{{{namespaces['w']}}}pict", f"{{{namespaces['mc']}}}AlternateContent", f"{{{namespaces['w']}}}t"] for child in r):
                                    if r.getparent() is not None:
                                        r.getparent().remove(r)
                    else:
                        # Simple 1:1 replacement in the primary run
                        target_idx = 0
                        # Inherit formatting from the run that was longest originally
                        lens = ["".join(t.text for t in ts if t.text) for r, ts in text_bearing_runs]
                        target_idx = lens.index(max(lens)) if lens else 0
                        
                        for i, (r, ts) in enumerate(text_bearing_runs):
                            if i == target_idx:
                                ts[0].text = sanitized_full
                                for t in ts[1:]: t.getparent().remove(t)
                            else:
                                for t in ts: 
                                    if t.getparent() == r: r.remove(t)
                                if not any(child.tag in [f"{{{namespaces['w']}}}drawing", f"{{{namespaces['w']}}}pict", f"{{{namespaces['mc']}}}AlternateContent", f"{{{namespaces['w']}}}t"] for child in r):
                                    if r.getparent() is not None:
                                        r.getparent().remove(r)
                
                modified_files.add(para_file)

            # HEADER SHADING RECT ADJUSTMENT:
            # After updating text, check if there is a shading rectangle covering the header
            # and resize it so it always fully covers the header text (handles multi-line titles).
            if self._adjust_header_shading_rect_height(root, namespaces, full_text_items):
                modified_files.add(document_xml_path)

            # Write all modified XML files back
            for file_path in modified_files:
                if file_path == document_xml_path:
                    tree.write(file_path, xml_declaration=True, encoding='UTF-8', standalone=True)
            
            # Repackage as DOCX by copying the ORIGINAL zip, replacing only modified files
            # This perfectly preserves Microsoft Word's specific compression methods and ZIP metadata
            with zipfile.ZipFile(original_file_path, 'r') as zin, zipfile.ZipFile(output_path, 'w') as zout:
                for item in zin.infolist():
                    # Windows zip paths use forward slash, but os.path.join handles normalization
                    temp_file_path = os.path.normpath(os.path.join(temp_dir, item.filename))
                    
                    # Normalize modified files for reliable checks
                    normalized_modified = {os.path.normpath(f) for f in modified_files}
                    
                    if temp_file_path in normalized_modified:
                        with open(temp_file_path, 'rb') as f:
                            zout.writestr(item, f.read())
                    else:
                        zout.writestr(item, zin.read(item.filename))
            
            print(f"✓ Successfully updated {len(all_paragraphs)} paragraphs with XML-level preservation")
    
    def _adjust_header_shading_rect_height(self, root, namespaces: dict, full_text_items: list) -> bool:
        """
        Detect a shading/background rectangle at the top of the document and adjust its height
        so it always fully covers the header text block.

        Strategy:
        - Look for VML rects (v:rect) or WordprocessingShape rectangles (wps:wsp) in the first
          few paragraphs of the body whose style contains a 'height' value.
        - Measure the actual heights of all header paragraphs (those that sit within the original
          shading rect's vertical span) by summing font-size × estimated line count × spacing.
        - Update the height in the shape's style attribute if the content requires more (or less)
          space than the original rectangle.

        Returns True if any modification was made.
        """
        try:
            W  = namespaces['w']
            V  = namespaces['v']
            MC = namespaces['mc']
            WPS = namespaces['wps']
            A   = namespaces['a']

            body = root.find('w:body', namespaces)
            if body is None:
                return False

            body_paragraphs = body.findall('w:p', namespaces)

            # ------------------------------------------------------------------ #
            # 1. Locate the shading rectangle and record its current height (pt). #
            # ------------------------------------------------------------------ #
            shape_elem = None          # The element whose style we will update
            shape_style_attr = None    # The attribute name that holds the style string
            original_height_pt = None  # Current height value in points

            # Search within the first 10 paragraphs — the rect is almost always in the header block
            for para in body_paragraphs[:10]:
                # Check both VML (v:rect / v:shape) and DrawingML (wps:wsp) shapes
                for elem in para.iter():
                    local = etree.QName(elem.tag).localname if '}' in elem.tag else elem.tag

                    if local in ('rect', 'shape', 'roundrect'):
                        style_val = elem.get('style', '')
                        if style_val and 'height' in style_val:
                            # Parse the height from the style string, e.g. "height:57pt" or "height:1.5in"
                            h_match = re.search(r'height\s*:\s*([\d.]+)(pt|in|cm|mm)?', style_val)
                            if h_match:
                                raw_val = float(h_match.group(1))
                                unit = (h_match.group(2) or 'pt').lower()
                                # Convert everything to points
                                if unit == 'in':
                                    raw_val *= 72
                                elif unit == 'cm':
                                    raw_val *= 28.3465
                                elif unit == 'mm':
                                    raw_val *= 2.83465
                                original_height_pt = raw_val
                                shape_elem = elem
                                shape_style_attr = 'style'
                                break

                    elif local == 'spPr':  # DrawingML shape properties
                        # Look for a:xfrm/a:ext with cy attribute (EMU units)
                        xfrm = elem.find(f'{{{A}}}xfrm', namespaces) if A in namespaces else None
                        if xfrm is None:
                            xfrm = elem.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}xfrm')
                        if xfrm is not None:
                            ext = xfrm.find('{http://schemas.openxmlformats.org/drawingml/2006/main}ext')
                            if ext is not None:
                                cy_emu = ext.get('cy')
                                if cy_emu:
                                    # 914400 EMU = 1 inch = 72 pt
                                    original_height_pt = int(cy_emu) * 72 / 914400
                                    shape_elem = ext
                                    shape_style_attr = 'cy'   # sentinel: means EMU attribute
                                    break

                if shape_elem is not None:
                    break

            if shape_elem is None or original_height_pt is None:
                print("   ℹ️  No header shading rectangle detected — skipping height adjustment.")
                return False

            print(f"   📐 Header shading rect found. Current height: {original_height_pt:.1f}pt")

            # ------------------------------------------------------------------ #
            # 2. Estimate the total height required to cover all header content.  #
            # ------------------------------------------------------------------ #
            # We consider the paragraphs whose cumulative estimated height does   #
            # not exceed the original rect height plus a reasonable overflow.     #
            # Page width inside margins is typically ~468pt (6.5") for letter.   #

            # Retrieve page/margin info for line-wrap estimation
            sect_pr = body.find('w:sectPr', namespaces)
            page_width_twips  = 12240  # default letter width
            margin_left_twips = 1440
            margin_right_twips = 1440
            if sect_pr is not None:
                pg_sz  = sect_pr.find('w:pgSz',  namespaces)
                pg_mar = sect_pr.find('w:pgMar', namespaces)
                if pg_sz  is not None: page_width_twips  = int(pg_sz.get(f'{{{W}}}w',  page_width_twips))
                if pg_mar is not None:
                    margin_left_twips  = int(pg_mar.get(f'{{{W}}}left',  margin_left_twips))
                    margin_right_twips = int(pg_mar.get(f'{{{W}}}right', margin_right_twips))

            # Usable text width in points (1 twip = 1/20 pt)
            usable_width_pt = (page_width_twips - margin_left_twips - margin_right_twips) / 20

            def _estimate_para_height_pt(para_elem, text_override=None) -> float:
                """Rough height estimate for a paragraph in points."""
                # Collect font size from the first text-bearing run's rPr, or pPr default
                font_size_pt = 11.0  # fallback

                # Check paragraph-level default
                pPr = para_elem.find('w:pPr', namespaces)
                if pPr is not None:
                    pStyle = pPr.find('w:pStyle', namespaces)
                    rPr_default = pPr.find('w:rPr', namespaces)
                    if rPr_default is not None:
                        sz = rPr_default.find('w:sz', namespaces)
                        if sz is not None:
                            val = sz.get(f'{{{W}}}val')
                            if val:
                                font_size_pt = int(val) / 2  # half-points

                # Check first text run's rPr
                for run in para_elem.findall('.//w:r', namespaces):
                    rPr = run.find('w:rPr', namespaces)
                    if rPr is not None:
                        sz = rPr.find('w:sz', namespaces)
                        if sz is not None:
                            val = sz.get(f'{{{W}}}val')
                            if val:
                                font_size_pt = int(val) / 2
                                break

                # Line spacing — check for w:spacing w:line
                line_spacing_factor = 1.15  # default
                if pPr is not None:
                    spacing = pPr.find('w:spacing', namespaces)
                    if spacing is not None:
                        line_rule = spacing.get(f'{{{W}}}lineRule', 'auto')
                        line_val  = spacing.get(f'{{{W}}}line')
                        if line_val and line_rule in ('auto', 'atLeast'):
                            # 240 = single spacing
                            line_spacing_factor = int(line_val) / 240

                # Space before/after in pt
                space_before_pt = 0.0
                space_after_pt  = 0.0
                if pPr is not None:
                    spacing = pPr.find('w:spacing', namespaces)
                    if spacing is not None:
                        sb = spacing.get(f'{{{W}}}before')
                        sa = spacing.get(f'{{{W}}}after')
                        if sb: space_before_pt = int(sb) / 20
                        if sa: space_after_pt  = int(sa) / 20

                # Paragraph text (use override if provided, e.g. the new AI-generated text)
                if text_override is not None:
                    text = text_override
                else:
                    text = ''.join(t.text for t in para_elem.findall('.//w:t', namespaces) if t.text)
                text = re.sub(r'\*\*', '', text)  # strip markdown bold markers

                if not text.strip():
                    return space_before_pt + space_after_pt  # empty paragraph = just spacing

                # Estimate characters per line based on font size and usable width
                # Average character width ≈ font_size_pt * 0.5 (rough heuristic for proportional fonts)
                avg_char_width_pt = font_size_pt * 0.5
                chars_per_line = max(1, int(usable_width_pt / avg_char_width_pt))

                # Count lines (word-wrap aware)
                words = text.split()
                lines = 1
                current_line_len = 0
                for word in words:
                    word_len = len(word) + 1  # +1 for space
                    if current_line_len + word_len > chars_per_line:
                        lines += 1
                        current_line_len = word_len
                    else:
                        current_line_len += word_len

                line_height_pt = font_size_pt * line_spacing_factor
                return space_before_pt + (lines * line_height_pt) + space_after_pt

            # Map AI text items back to paragraphs so we can use the new text for estimates
            # all_paragraphs list is out of scope here, so we rebuild a para→new_text map
            # by walking body paragraphs and matching them in order with full_text_items
            ai_text_by_para = {}  # id(para_elem) -> new text string
            text_iter = iter(full_text_items)
            for para in body_paragraphs:
                t_content = ''.join(t.text for t in para.findall('.//w:t', namespaces) if t.text)
                if t_content.strip():
                    new_t = next(text_iter, None)
                    if new_t is not None:
                        ai_text_by_para[id(para)] = str(new_t)

            # Walk header paragraphs until cumulative estimated height exceeds
            # 2× the original rect height (a generous upper bound search window)
            cumulative_pt = 0.0
            header_para_count = 0
            search_limit_pt = original_height_pt * 2.5

            for para in body_paragraphs:
                est = _estimate_para_height_pt(
                    para,
                    text_override=ai_text_by_para.get(id(para))
                )
                cumulative_pt += est
                header_para_count += 1

                # Check if this paragraph has any non-header indicator (horizontal rule,
                # long cap-letter section title, or if we've already gone well past original rect)
                if cumulative_pt > search_limit_pt:
                    break

                # Stop when we've accounted for more than the original rect height
                # — anything beyond that is body content, not header
                if cumulative_pt >= original_height_pt:
                    break

            needed_height_pt = cumulative_pt
            print(f"   📏 Estimated header content height: {needed_height_pt:.1f}pt (original rect: {original_height_pt:.1f}pt)")

            # Add a small padding buffer beneath the last header line
            PADDING_PT = 6.0
            needed_height_pt += PADDING_PT

            # Only update if the difference is significant (> 2pt) to avoid spurious changes
            if abs(needed_height_pt - original_height_pt) < 2.0:
                print("   ✅ Header shading rect height is already correct — no change needed.")
                return False

            # ------------------------------------------------------------------ #
            # 3. Apply the new height to the shape element.                       #
            # ------------------------------------------------------------------ #
            if shape_style_attr == 'cy':
                # DrawingML: update the 'cy' attribute in EMU
                new_cy_emu = int(needed_height_pt * 914400 / 72)
                shape_elem.set('cy', str(new_cy_emu))
                print(f"   ✏️  Updated DrawingML shape cy: {new_cy_emu} EMU ({needed_height_pt:.1f}pt)")
            else:
                # VML: update the height in the style string
                current_style = shape_elem.get('style', '')
                # Replace height:XXXpt (or height:X.Xin etc.) with new value in pt
                new_height_str = f'{needed_height_pt:.2f}pt'
                new_style = re.sub(
                    r'height\s*:\s*[\d.]+(?:pt|in|cm|mm)?',
                    f'height:{new_height_str}',
                    current_style
                )
                shape_elem.set('style', new_style)
                print(f"   ✏️  Updated VML shape height: {new_height_str}")

            return True

        except Exception as e:
            print(f"   ⚠️  Error adjusting header shading rect height: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _adjust_margins_for_page_count(self, root, namespaces: dict, document_xml_path: str, 
                                        original_file_path: str, temp_dir: str, modified_files: set):
        """
        Adjust margins intelligently to try to keep the same page count as the original.
        
        Rules:
        1. First adjust top/bottom margins equally (minimum 0.5" = 720 twips)
        2. Then adjust left/right margins equally (minimum 0.5" = 720 twips)
        3. Always keep top=bottom and left=right
        
        Note: 1 inch = 1440 twips, 0.5 inch = 720 twips
        """
        try:
            # Estimate original page count
            original_pages = self.estimate_page_count_from_content(original_file_path)
            print(f"   📄 Original document: ~{original_pages} page(s)")
            
            # Get section properties
            sect_pr = root.find('.//w:body/w:sectPr', namespaces)
            if sect_pr is None:
                print("   ⚠️  No section properties found, skipping margin adjustment")
                return
            
            pg_mar = sect_pr.find('w:pgMar', namespaces)
            if pg_mar is None:
                print("   ⚠️  No page margins found, skipping margin adjustment")
                return
            
            # Get current margins (default 1440 twips = 1 inch)
            top = int(pg_mar.get(f'{{{namespaces["w"]}}}top', 1440))
            bottom = int(pg_mar.get(f'{{{namespaces["w"]}}}bottom', 1440))
            left = int(pg_mar.get(f'{{{namespaces["w"]}}}left', 1440))
            right = int(pg_mar.get(f'{{{namespaces["w"]}}}right', 1440))
            
            MIN_MARGIN = 720  # 0.5 inches
            STEP = 72  # 0.05 inches per step
            
            # Current margins
            current_top_bottom = min(top, bottom)
            current_left_right = min(left, right)
            
            print(f"   📏 Current margins: T/B={current_top_bottom/1440:.2f}\", L/R={current_left_right/1440:.2f}\"")
            
            # Strategy: Aggressive margin reduction for PDF conversion
            # LibreOffice PDF rendering adds extra spacing, so we need tighter margins
            # to ensure the PDF matches the page count
            # Step 1: Reduce top/bottom first (up to minimum of 0.5")
            # Step 2: Reduce left/right if needed (up to minimum of 0.5")
            
            # First, reduce top/bottom margins more aggressively
            new_top_bottom = current_top_bottom
            if current_top_bottom > MIN_MARGIN:
                # Reduce by up to 0.35 inches (504 twips) - more aggressive for PDF
                # This compensates for LibreOffice's PDF renderer adding extra line spacing
                reduction = min(504, current_top_bottom - MIN_MARGIN)
                new_top_bottom = current_top_bottom - reduction
                new_top_bottom = max(MIN_MARGIN, new_top_bottom)
            
            # Also reduce left/right margins if needed
            new_left_right = current_left_right
            if current_left_right > MIN_MARGIN:
                # Reduce by up to 0.25 inches (360 twips)
                reduction = min(360, current_left_right - MIN_MARGIN)
                new_left_right = current_left_right - reduction
                new_left_right = max(MIN_MARGIN, new_left_right)
            
            # Apply new margins (keep them equal)
            pg_mar.set(f'{{{namespaces["w"]}}}top', str(new_top_bottom))
            pg_mar.set(f'{{{namespaces["w"]}}}bottom', str(new_top_bottom))
            pg_mar.set(f'{{{namespaces["w"]}}}left', str(new_left_right))
            pg_mar.set(f'{{{namespaces["w"]}}}right', str(new_left_right))
            
            print(f"   ✨ Adjusted margins: T/B={new_top_bottom/1440:.2f}\", L/R={new_left_right/1440:.2f}\" (targeting {original_pages} page(s))")
            modified_files.add(document_xml_path)
            
        except Exception as e:
            print(f"   ⚠️  Error adjusting margins: {e}")
    
    def create_docx_preserve_formatting(self, original_file_path: str, resume_data: Dict[str, Any], output_path: str):
        """
        Create a DOCX file by copying the original and updating text paragraph-by-paragraph.
        This preserves ALL formatting: fonts, colors, spacing, styles, alignment, etc.
        """
        import shutil
        
        # Copy the original file to the output path first
        shutil.copy2(original_file_path, output_path)
        
        # Now open the copied document
        doc = Document(output_path)
        
        # Get full text from AI-generated resume in order
        full_text_items = resume_data.get("full_text", [])
        
        if not full_text_items:
            print("⚠️  Warning: No full_text in resume_data, falling back to sections")
            # Fallback: build from sections
            full_text_items = []
            for section in resume_data.get("sections", []):
                if section.get("type") == "table":
                    continue
                if section.get("title"):
                    full_text_items.append(section["title"])
                for item in section.get("content", []):
                    if item:
                        full_text_items.append(str(item))
        
        # Get all non-empty paragraphs from the original
        original_paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                original_paragraphs.append(para)
        
        # Check if counts match
        if len(original_paragraphs) != len(full_text_items):
            print(f"⚠️  Warning: Paragraph count mismatch!")
            print(f"   Original: {len(original_paragraphs)} paragraphs")
            print(f"   AI Generated: {len(full_text_items)} text items")
            print(f"   Using fallback method.")
            # Recreate fresh document
            self.create_docx(resume_data, output_path)
            return
        
        # Update each paragraph's text while keeping its formatting
        for para, new_text in zip(original_paragraphs, full_text_items):
            # Sanitize the new text
            sanitized_text = self.sanitize_text(str(new_text))
            
            # Check if the paragraph has a pattern like "Bold: Normal" (e.g., Core Competencies)
            # This is indicated by multiple runs or a colon in the text
            has_colon_pattern = ':' in sanitized_text and len(para.runs) > 1
            
            if has_colon_pattern:
                # Handle multi-run formatting (e.g., "Leadership & Strategy: IT Roadmap Development...")
                # Split at the colon to preserve bold/normal pattern
                parts = sanitized_text.split(':', 1)
                if len(parts) == 2:
                    label_part = parts[0] + ':'  # Include the colon in the bold part
                    content_part = parts[1]
                    
                    # Check if first run is bold
                    first_run_bold = para.runs[0].bold if para.runs[0].bold is not None else False
                    
                    if first_run_bold:
                        # Clear all runs
                        for run in para.runs:
                            run.text = ""
                        
                        # Use first run for bold label
                        para.runs[0].text = label_part
                        
                        # Add second run for normal content if we have multiple runs
                        if len(para.runs) > 1:
                            para.runs[1].text = content_part
                        else:
                            # Create a new run with normal (non-bold) formatting
                            new_run = para.add_run(content_part)
                            if para.runs[0].font.name:
                                new_run.font.name = para.runs[0].font.name
                            if para.runs[0].font.size:
                                new_run.font.size = para.runs[0].font.size
                            new_run.bold = False
                    else:
                        # Not a bold pattern, treat normally
                        para.runs[0].text = sanitized_text
                        for run in para.runs[1:]:
                            run.text = ""
                else:
                    # Couldn't split properly, use default behavior
                    para.runs[0].text = sanitized_text
                    for run in para.runs[1:]:
                        run.text = ""
            else:
                # Single-run paragraph or no colon pattern
                if para.runs:
                    # Update the first run with all the text
                    para.runs[0].text = sanitized_text
                    # Clear other runs
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    # No runs, just set paragraph text
                    para.text = sanitized_text
        
        # Save the document
        doc.save(output_path)
        print(f"✓ Successfully updated {len(original_paragraphs)} paragraphs with preserved formatting")
    def create_redline_docx(self, original_file_path: str, resume_data: Dict[str, Any], output_path: str):
        """
        Create a 'Redline' DOCX where changed text is highlighted in Red.
        Uses granular diffs to highlight specific words.
        """
        import shutil
        import difflib
        from docx import Document
        from docx.shared import RGBColor
        
        # Copy the original file to the output path first
        shutil.copy2(original_file_path, output_path)
        doc = Document(output_path)
        
        full_text_items = resume_data.get("full_text", [])
        if not full_text_items: return
            
        original_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    
        if len(original_paragraphs) != len(full_text_items):
            print(f"⚠️  Redline count mismatch: {len(original_paragraphs)} vs {len(full_text_items)}. Proceeding with partial diff.")
    
        for para, new_text in zip(original_paragraphs, full_text_items):
            original_text = para.text
            sanitized_new = self.sanitize_text(str(new_text))
            
            # Helper to rebuild paragraph with diffs
            def build_diff_runs(target_para, old_txt, new_txt, bold=False):
                # Diff logic
                # We split by words for better readability than char-level
                a_words = old_txt.split(' ')
                b_words = new_txt.split(' ')
                
                matcher = difflib.SequenceMatcher(None, a_words, b_words)
                
                for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                    chunk_text = " ".join(b_words[j1:j2]) + " "
                    
                    if not chunk_text.strip(): continue
                        
                    if tag == 'equal':
                        run = target_para.add_run(chunk_text)
                        run.bold = bold
                    elif tag in ('replace', 'insert'):
                        run = target_para.add_run(chunk_text)
                        run.bold = bold
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    # delete is ignored
            
            # Check patterns
            # Identify if this looks like a bold header line (e.g. "Key Achievement: ...")
            # We use a heuristic or check existing runs
            has_colon_pattern = ':' in sanitized_new and len(para.runs) > 1
            
            # Clear existing content
            para.text = "" 
            
            if has_colon_pattern:
                parts = sanitized_new.split(':', 1)
                orig_parts = original_text.split(':', 1)
                
                if len(parts) == 2:
                    label_part = parts[0] + ':'
                    content_part = parts[1].strip()
                    
                    # Add label (assume BOLD)
                    run = para.add_run(label_part + " ")
                    run.bold = True
                    
                    # Diff the content part
                    orig_content = orig_parts[1] if len(orig_parts) == 2 else ""
                    build_diff_runs(para, orig_content, content_part, bold=False)
                else:
                    build_diff_runs(para, original_text, sanitized_new)
            else:
                build_diff_runs(para, original_text, sanitized_new)
        
        doc.save(output_path)

    def create_docx(self, resume_data: Dict[str, Any], output_path: str):
        """
        Create a DOCX file from structured resume data.
        """

        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Process sections
        sections_data = resume_data.get("sections", [])
        
        for section in sections_data:
            title = section.get("title", "")
            content = section.get("content", [])
            section_type = section.get("type", "text")
            
            if section_type == "table":
                # Handle tables
                if isinstance(content, list) and content:
                    table = doc.add_table(rows=len(content), cols=len(content[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for i, row_data in enumerate(content):
                        row = table.rows[i]
                        for j, cell_text in enumerate(row_data):
                            row.cells[j].text = self.sanitize_text(str(cell_text))
            else:
                # Add section title
                if title:
                    heading = doc.add_paragraph(self.sanitize_text(title))
                    heading.runs[0].bold = True
                    heading.runs[0].font.size = Pt(14)
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                
                # Add section content
                for item in content:
                    if item:
                        para = doc.add_paragraph(self.sanitize_text(str(item)))
                        para.runs[0].font.size = Pt(11)
        
        doc.save(output_path)
    
    def create_pdf_from_docx(self, docx_path: str, output_path: str) -> dict:
        """
        Convert a DOCX file to PDF using the best available method.
        
        Priority:
        1. LibreOffice (best - perfect DOCX formatting preservation)
        2. DOCX → HTML → PDF (good for basic documents)
        3. Pandoc with CSS-based engines (fallback)
        4. Pure Python reportlab (basic fallback)
        
        Returns dict with success status and font information.
        """
        # Check for missing fonts before conversion
        font_info = {'fonts_detected': [], 'missing_fonts': [], 'all_available': True}
        try:
            from services.font_service import FontService
            font_service = FontService()
            font_info = font_service.ensure_fonts_available(docx_path)
        except Exception as e:
            print(f"   ℹ️  Font check skipped: {e}")
        
        # Try LibreOffice FIRST - best for preserving DOCX formatting
        success = False
        if self._try_libreoffice_conversion(docx_path, output_path):
            success = True
        elif self._try_html_intermediate_conversion(docx_path, output_path):
            success = True
        elif self._try_pandoc_conversion(docx_path, output_path):
            success = True
        else:
            # Last resort: Pure Python implementation
            print("   ℹ️  No converters available, using basic reportlab fallback...")
            success = self._convert_with_reportlab(docx_path, output_path)
        
        return {
            'success': success,
            'font_info': font_info
        }
    
    def _try_html_intermediate_conversion(self, docx_path: str, output_path: str) -> bool:
        """
        Convert DOCX → HTML → PDF for better formatting preservation.
        Uses Pandoc for DOCX→HTML and WeasyPrint for HTML→PDF.
        """
        import tempfile
        
        try:
            # Setup environment to include venv binaries
            env = os.environ.copy()
            venv_bin = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'venv', 'bin')
            if os.path.exists(venv_bin):
                env['PATH'] = f"{venv_bin}:{env.get('PATH', '')}"
            
            # Check if pandoc is available
            result = subprocess.run(
                ['pandoc', '--version'],
                capture_output=True,
                timeout=5,
                text=True,
                env=env
            )
            
            if result.returncode != 0:
                return False
            
            # Check if weasyprint is available
            result = subprocess.run(
                ['weasyprint', '--version'],
                capture_output=True,
                timeout=5,
                text=True,
                env=env
            )
            
            if result.returncode != 0:
                print("   ℹ️  WeasyPrint not found for HTML conversion")
                return False
            
            # Step 1: Convert DOCX to HTML with embedded styles
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False) as html_file:
                html_path = html_file.name
            
            try:
                # Convert DOCX to standalone HTML with CSS
                pandoc_cmd = [
                    'pandoc',
                    docx_path,
                    '-o', html_path,
                    '--standalone',
                    '--self-contained',
                    '--css', 'style.css'  # Pandoc will embed default styles
                ]
                
                result = subprocess.run(
                    pandoc_cmd,
                    capture_output=True,
                    timeout=30,
                    text=True,
                    env=env
                )
                
                if result.returncode != 0 or not os.path.exists(html_path):
                    print(f"   ℹ️  DOCX→HTML conversion failed")
                    return False
                
                # Step 2: Convert HTML to PDF using WeasyPrint
                weasyprint_cmd = [
                    'weasyprint',
                    html_path,
                    output_path
                ]
                
                result = subprocess.run(
                    weasyprint_cmd,
                    capture_output=True,
                    timeout=30,
                    text=True,
                    env=env
                )
                
                if result.returncode == 0 and os.path.exists(output_path):
                    print(f"   ✓ PDF generated via HTML (DOCX→HTML→PDF)")
                    return True
                else:
                    print(f"   ℹ️  HTML→PDF conversion failed: {result.stderr[:200]}")
                    return False
                    
            finally:
                # Clean up temporary HTML file
                if os.path.exists(html_path):
                    os.unlink(html_path)
                    
        except subprocess.TimeoutExpired:
            print("   ⚠️  HTML conversion timed out")
            return False
        except Exception as e:
            print(f"   ℹ️  HTML intermediate conversion error: {e}")
            return False
    
    def _try_libreoffice_conversion(self, docx_path: str, output_path: str) -> bool:
        """
        Try to convert DOCX to PDF using LibreOffice.
        This preserves the most formatting including shapes, colors, and headers.
        """
        try:
            # Check for LibreOffice
            libreoffice_cmd = None
            for cmd in ['libreoffice', 'soffice', '/usr/bin/libreoffice', '/usr/bin/soffice']:
                try:
                    result = subprocess.run(
                        [cmd, '--version'],
                        capture_output=True,
                        timeout=5,
                        text=True
                    )
                    if result.returncode == 0:
                        libreoffice_cmd = cmd
                        break
                except (FileNotFoundError, subprocess.TimeoutExpired):
                    continue
            
            if not libreoffice_cmd:
                return False
            
            # Get absolute paths
            output_dir = os.path.dirname(os.path.abspath(output_path))
            docx_abs_path = os.path.abspath(docx_path)
            
            # Convert with proper PDF export filter to preserve formatting
            # Using filter:writer_pdf_Export with specific options
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'pdf:writer_pdf_Export',
                '--outdir', output_dir,
                docx_abs_path
            ]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                timeout=60,
                text=True
            )
            
            if result.returncode != 0:
                print(f"   ℹ️  LibreOffice conversion failed: {result.stderr[:200]}")
                return False
            
            # LibreOffice creates PDF with same base name
            base_name = os.path.splitext(os.path.basename(docx_path))[0]
            generated_pdf = os.path.join(output_dir, f"{base_name}.pdf")
            
            # Rename if needed
            if os.path.exists(generated_pdf) and generated_pdf != output_path:
                shutil.move(generated_pdf, output_path)
            
            if os.path.exists(output_path):
                print(f"   ✓ PDF generated with LibreOffice (formatting preserved)")
                return True
            
            return False
                
        except subprocess.TimeoutExpired:
            print("   ⚠️  LibreOffice conversion timed out")
            return False
        except Exception as e:
            print(f"   ℹ️  LibreOffice error: {e}")
            return False
    
    def _try_pandoc_conversion(self, docx_path: str, output_path: str) -> bool:
        """
        Try to convert DOCX to PDF using Pandoc with CSS-based engines.
        CSS-based engines preserve layout better than LaTeX.
        
        Priority:
        1. weasyprint (CSS-based, good for layout preservation)
        2. prince (CSS-based, commercial but high quality)
        3. wkhtmltopdf (WebKit-based, decent fallback)
        """
        try:
            # Setup environment to include venv binaries
            env = os.environ.copy()
            venv_bin = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'venv', 'bin')
            if os.path.exists(venv_bin):
                env['PATH'] = f"{venv_bin}:{env.get('PATH', '')}"
            
            # Check if pandoc is available
            result = subprocess.run(
                ['pandoc', '--version'],
                capture_output=True,
                timeout=5,
                text=True,
                env=env
            )
            
            if result.returncode != 0:
                return False
            
            # Try different PDF engines in order of preference
            pdf_engines = [
                ('weasyprint', 'WeasyPrint (CSS-based)'),
                ('prince', 'Prince (CSS-based)'),
                ('wkhtmltopdf', 'wkhtmltopdf'),
            ]
            
            for engine, engine_name in pdf_engines:
                try:
                    # Check if the engine is available
                    engine_check = subprocess.run(
                        [engine, '--version'] if engine != 'prince' else ['prince', '--version'],
                        capture_output=True,
                        timeout=5,
                        text=True,
                        stderr=subprocess.DEVNULL,
                        env=env
                    )
                    
                    if engine_check.returncode != 0:
                        continue
                    
                    # Try conversion with this engine
                    cmd = [
                        'pandoc',
                        docx_path,
                        '-o', output_path,
                        f'--pdf-engine={engine}'
                    ]
                    
                    result = subprocess.run(
                        cmd,
                        capture_output=True,
                        timeout=60,
                        text=True,
                        env=env
                    )
                    
                    if result.returncode == 0 and os.path.exists(output_path):
                        print(f"   ✓ PDF generated with Pandoc using {engine_name}")
                        return True
                    
                except (FileNotFoundError, subprocess.TimeoutExpired):
                    continue
            
            # If no CSS-based engine worked, inform user
            print(f"   ℹ️  Pandoc available but no suitable PDF engine found")
            print(f"   💡 Install weasyprint: pip install weasyprint")
            return False
                
        except (FileNotFoundError, subprocess.TimeoutExpired):
            return False
        except Exception as e:
            print(f"   ℹ️  Pandoc error: {e}")
            return False
    
    def _convert_with_reportlab(self, docx_path: str, output_path: str) -> bool:
        """
        Convert DOCX to PDF using pure Python (reportlab).
        Fallback method when Pandoc is not available.
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.units import inch
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
            from reportlab.lib import colors
            from docx import Document
            from docx.shared import Pt, Inches as DocxInches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Open the DOCX file
            doc = Document(docx_path)
            
            # Get page setup from DOCX
            section = doc.sections[0]
            page_width = float(section.page_width) / 914400 * inch  # Convert EMUs to inches
            page_height = float(section.page_height) / 914400 * inch
            top_margin = float(section.top_margin) / 914400 * inch
            bottom_margin = float(section.bottom_margin) / 914400 * inch
            left_margin = float(section.left_margin) / 914400 * inch
            right_margin = float(section.right_margin) / 914400 * inch
            
            # Create PDF with matching page setup
            pdf = SimpleDocTemplate(
                output_path,
                pagesize=(page_width, page_height),
                topMargin=top_margin,
                bottomMargin=bottom_margin,
                leftMargin=left_margin,
                rightMargin=right_margin
            )
            
            # Build styles
            styles = getSampleStyleSheet()
            story = []
            
            # Process paragraphs
            for para in doc.paragraphs:
                if not para.text.strip():
                    story.append(Spacer(1, 0.1*inch))
                    continue
                
                # Determine alignment
                alignment = TA_LEFT
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    alignment = TA_CENTER
                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    alignment = TA_RIGHT
                elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    alignment = TA_JUSTIFY
                
                # Get formatting from first run
                font_size = 11
                font_name = 'Helvetica'
                is_bold = False
                is_italic = False
                text_color = colors.black
                
                # Standard PDF fonts supported by reportlab natively
                standard_fonts = ['Courier', 'Helvetica', 'Times-Roman', 'Symbol', 'ZapfDingbats']
                
                if para.runs:
                    first_run = para.runs[0]
                    if first_run.font.size:
                        font_size = first_run.font.size.pt
                    if first_run.font.name:
                        # Reportlab crashes on non-standard fonts without explicit TTF registration
                        doc_font = first_run.font.name
                        if any(s in doc_font for s in standard_fonts):
                            font_name = doc_font
                        else:
                            font_name = 'Helvetica'
                    is_bold = first_run.bold if first_run.bold is not None else False
                    is_italic = first_run.italic if first_run.italic is not None else False
                    
                    # Get color
                    if first_run.font.color and first_run.font.color.rgb:
                        rgb = first_run.font.color.rgb
                        text_color = colors.Color(rgb[0]/255, rgb[1]/255, rgb[2]/255)
                
                # Create paragraph style
                style = ParagraphStyle(
                    'CustomStyle',
                    parent=styles['Normal'],
                    fontSize=font_size,
                    fontName=font_name,
                    alignment=alignment,
                    textColor=text_color,
                    spaceAfter=6
                )
                
                if is_bold:
                    style.fontName = font_name + '-Bold' if font_name == 'Helvetica' else font_name
                
                # Build formatted text with run-level formatting
                formatted_text = ""
                from docx.text.run import Run
                
                # To include hyperlinks which python-docx skips in para.runs
                for child in para._element:
                    runs_to_process = []
                    
                    if child.tag.endswith('}r'):
                        runs_to_process.append(Run(child, para))
                    elif child.tag.endswith('}hyperlink'):
                        for r_node in child.findall('.//w:r', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            runs_to_process.append(Run(r_node, para))
                            
                    for run in runs_to_process:
                        text = self.sanitize_text(run.text)
                        if not text:
                            continue
                            
                        # Build formatting tags
                        is_run_bold = run.bold if run.bold is not None else False
                        is_run_italic = run.italic if run.italic is not None else False
                        is_run_underline = run.underline if run.underline is not None else False
                        
                        # Apply fallback formatting if we detected a hyperlink run
                        # (hyperlink runs often lose explicit visual formatting info when instantiated dynamically)
                        if child.tag.endswith('}hyperlink'):
                            # Hyperlinks are usually blue and underlined
                            is_run_underline = True
                            text = f'<font color="blue">{text}</font>'
                            
                        if is_run_bold:
                            text = f"<b>{text}</b>"
                        if is_run_italic:
                            text = f"<i>{text}</i>"
                        if is_run_underline:
                            text = f"<u>{text}</u>"
                        
                        formatted_text += text
                
                if formatted_text:
                    try:
                        p = Paragraph(formatted_text, style)
                        story.append(p)
                    except Exception as e:
                        # Fallback to plain text if formatting fails
                        p = Paragraph(self.sanitize_text(para.text), style)
                        story.append(p)
            
            # Process tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    t = Table(table_data)
                    t.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('PADDING', (0, 0), (-1, -1), 6),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 0.2*inch))
            
            # Build PDF
            pdf.build(story)
            print(f"   ✓ PDF generated successfully (pure Python)")
            return True
            
        except Exception as e:
            print(f"   ⚠️  Error converting DOCX to PDF: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def create_pdf(self, resume_data: Dict[str, Any], output_path: str):
        """
        DEPRECATED: Create a PDF file from structured resume data.
        This method is kept for backward compatibility but should not be used
        as it doesn't preserve DOCX formatting.
        Use create_pdf_from_docx() instead.
        """
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch
        )
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#000000',
            spaceAfter=12,
            bold=True
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6
        )
        
        story = []
        
        sections_data = resume_data.get("sections", [])
        
        for section in sections_data:
            title = section.get("title", "")
            content = section.get("content", [])
            section_type = section.get("type", "text")
            
            if section_type != "table":
                # Add section title
                if title:
                    story.append(Paragraph(self.sanitize_text(title), heading_style))
                    story.append(Spacer(1, 0.1*inch))
                
                # Add section content
                for item in content:
                    if item:
                        story.append(Paragraph(self.sanitize_text(str(item)), body_style))
        
        doc.build(story)
    
    def create_txt(self, resume_data: Dict[str, Any], output_path: str):
        """
        Create a plain text file from structured resume data.
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            sections_data = resume_data.get("sections", [])
            
            for section in sections_data:
                title = section.get("title", "")
                content = section.get("content", [])
                section_type = section.get("type", "text")
                
                if section_type != "table":
                    # Add section title
                    if title:
                        sanitized_title = self.sanitize_text(title)
                        f.write(f"\n{sanitized_title.upper()}\n")
                        f.write("=" * len(sanitized_title) + "\n\n")
                    
                    # Add section content
                    for item in content:
                        if item:
                            f.write(f"{self.sanitize_text(str(item))}\n")
                    
                    f.write("\n")

    def create_cover_letter_docx(self, content: str, output_path: str):
        """Creates a simple formatted cover letter DOCX."""
        doc = Document()
        
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        paragraphs = content.split('\n')
        for p_text in paragraphs:
            p = doc.add_paragraph(p_text.strip())
            p.paragraph_format.space_after = Pt(6)
        
        doc.save(output_path)

    def extract_text(self, file_path: str) -> str:
        """Extract text from a file based on its extension."""
        if not file_path:
            return ""
        
        path_lower = file_path.lower()
        if path_lower.endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif path_lower.endswith('.txt'):
            return self.extract_text_from_txt(file_path)
        elif path_lower.endswith('.docx'):
            try:
                doc_data = self.parse_docx(file_path)
                return "\n".join(doc_data.get("full_text", []))
            except Exception as e:
                print(f"Error extracting text from docx: {e}")
                return ""
        else:
            return ""


    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using pdftotext."""
        try:
            result = subprocess.run(
                ['pdftotext', '-layout', pdf_path, '-'],
                capture_output=True,
                text=True,
                check=True
            )
            return result.stdout
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            return ""

    def extract_text_from_txt(self, txt_path: str) -> str:
        """Extract text from a plain text file."""
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading txt file: {e}")
            return ""
